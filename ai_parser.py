# ai_parser.py — AI-разбор ТЗ/сметы и сопоставление со справочником
import io, json, difflib, re
import streamlit as st


# ─── Извлечение текста из файла ───────────────────────────────────────────────

def extract_text_from_xlsx(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sh in wb.sheetnames:
        ws = wb[sh]
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if row_vals:
                lines.append("\t".join(row_vals))
    return "\n".join(lines)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]
    if ext in ("xlsx", "xls"):
        return extract_text_from_xlsx(file_bytes)
    elif ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    return ""


# ─── Вызов Claude API ─────────────────────────────────────────────────────────

SYSTEM_PROMPT = """Ты — эксперт по строительным сметам и ТЗ.
Твоя задача: извлечь из документа список строительных работ с единицами измерения и объёмами.
Верни ТОЛЬКО корректный JSON-массив без каких-либо пояснений и markdown.
Формат каждого объекта:
{"name": "Название работы на русском", "unit": "м²", "qty": 25.5}
Если объём не указан — ставь null. Если единица не указана — ставь "шт."
Не дублируй одинаковые позиции. Исключай итоговые строки, заголовки, реквизиты."""


def call_claude_api(text: str, api_key: str) -> list[dict]:
    """Вызывает Claude Haiku для разбора текста документа. Возвращает список {name, unit, qty}."""
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    # Обрезаем текст до 15000 символов чтобы не превысить контекст
    text_trimmed = text[:15000]

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4000,
        messages=[
            {
                "role": "user",
                "content": f"Извлеки список работ из следующего документа:\n\n{text_trimmed}"
            }
        ],
        system=SYSTEM_PROMPT,
    )

    raw = message.content[0].text.strip()
    # Убираем возможные ```json``` обёртки
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


# ─── Матчинг позиций справочника ──────────────────────────────────────────────

def match_items(parsed: list[dict], all_items: list) -> list[dict]:
    """
    Для каждой извлечённой позиции находит лучшее совпадение в справочнике.
    Возвращает список:
      {parsed_name, parsed_unit, qty, matched_item (или None), score}
    """
    names_lower = [i["name"].lower() for i in all_items]
    results = []
    for p in parsed:
        name = p.get("name", "")
        unit = p.get("unit", "")
        qty = p.get("qty")
        name_low = name.lower()

        # Сначала пробуем точное включение
        matches = difflib.get_close_matches(name_low, names_lower, n=3, cutoff=0.35)
        best_item = None
        best_score = 0.0
        if matches:
            best_name = matches[0]
            idx = names_lower.index(best_name)
            best_item = all_items[idx]
            best_score = difflib.SequenceMatcher(None, name_low, best_name).ratio()

        results.append({
            "parsed_name": name,
            "parsed_unit": unit,
            "qty": qty,
            "matched_item": best_item,
            "score": round(best_score, 2),
        })
    return results
